import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')  # 不使用Tcl/Tk后端，改用Agg后端
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签
plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号

def analyze_scores_by_class(excel_file, word_file="班级成绩分析报告.docx"):
    # 读取Excel文件
    df = pd.read_excel(excel_file)
    
    # 自动获取所有学科列名（排除'班级'、'序号'、'姓名'列）
    exclude_cols = ['班级', '序号', '姓名']
    subjects = [col for col in df.columns if col not in exclude_cols]
    
    # 定义分数段
    score_ranges = [
        (0, 60, '不及格'),
        (60, 70, '及格'),
        (70, 80, '中等'),
        (80, 90, '良好'),
        (90, 100, '优秀')
    ]
    
    # 创建Word文档
    doc = Document()
    doc.add_heading('班级成绩分析报告', 0)
    
    # 按班级分组计算统计数据
    class_groups = df.groupby('班级')
    
    for class_name, class_data in class_groups:
        doc.add_heading(f'{class_name}班级分析', level=1)
        
        # 为每个科目创建统计数据
        for subject in subjects:
            doc.add_heading(f'{subject}科目分析', level=2)
            
            # 计算该班级该科目的统计数据
            mean_score = class_data[subject].mean()
            pass_rate = (class_data[subject] >= 60).mean() * 100
            excellent_rate = (class_data[subject] >= 90).mean() * 100
            fail_rate = (class_data[subject] < 60).mean() * 100
            std_dev = class_data[subject].std()
            
            # 计算分数段人数统计
            score_stats = []
            total_students = len(class_data)
            for low, high, label in score_ranges:
                count = len(class_data[(class_data[subject] >= low) & (class_data[subject] < high)])
                percentage = (count / total_students) * 100
                score_stats.append((label, count, percentage))
            
            # 添加统计数据到Word
            doc.add_heading('基本统计数据', level=3)
            stats_table = doc.add_table(rows=5, cols=2)
            stats_table.style = 'Table Grid'
            
            stats = [
                ("平均分", f"{mean_score:.2f}"),
                ("及格率", f"{pass_rate:.2f}%"),
                ("优秀率", f"{excellent_rate:.2f}%"),
                ("差分率", f"{fail_rate:.2f}%"),
                ("标准差", f"{std_dev:.2f}")
            ]
            
            for i, (name, value) in enumerate(stats):
                stats_table.cell(i, 0).text = name
                stats_table.cell(i, 1).text = value
            
            # 添加分数段统计表格
            doc.add_heading('分数段统计', level=3)
            range_table = doc.add_table(rows=len(score_stats)+1, cols=3)
            range_table.style = 'Table Grid'
            
            # 添加表头
            headers = range_table.rows[0].cells
            headers[0].text = "分数段"
            headers[1].text = "人数"
            headers[2].text = "占比"
            
            # 填充数据
            for i, (label, count, percentage) in enumerate(score_stats, 1):
                row = range_table.rows[i].cells
                row[0].text = label
                row[1].text = str(count)
                row[2].text = f"{percentage:.2f}%"
            
            try:
                doc.add_heading('统计图表', level=3)
                
                # 1. 成绩分布直方图
                plt.figure(figsize=(10, 6))
                plt.hist(class_data[subject], bins=20, edgecolor='black')
                plt.title(f'{class_name}班级{subject}成绩分布直方图')
                plt.xlabel('分数')
                plt.ylabel('频数')
                plt.tight_layout()
                hist_file = f'hist_{class_name}_{subject}.png'
                plt.savefig(hist_file)
                plt.close()
                doc.add_picture(hist_file, width=Inches(6))
                
                # 2. 成绩箱线图
                plt.figure(figsize=(10, 6))
                plt.boxplot(class_data[subject])
                plt.title(f'{class_name}班级{subject}成绩箱线图')
                plt.ylabel('分数')
                plt.tight_layout()
                box_file = f'box_{class_name}_{subject}.png'
                plt.savefig(box_file)
                plt.close()
                doc.add_picture(box_file, width=Inches(6))
                
                # 3. 分数段人数比例饼图
                plt.figure(figsize=(10, 6))
                labels = [stat[0] for stat in score_stats]
                sizes = [stat[2] for stat in score_stats]
                plt.pie(sizes, labels=labels, autopct='%1.1f%%')
                plt.title(f'{class_name}班级{subject}分数段分布')
                plt.tight_layout()
                pie_file = f'pie_{class_name}_{subject}.png'
                plt.savefig(pie_file)
                plt.close()
                doc.add_picture(pie_file, width=Inches(6))
                
                # 4. 分数段人数柱状图
                plt.figure(figsize=(10, 6))
                plt.bar(labels, [stat[1] for stat in score_stats])
                plt.title(f'{class_name}班级{subject}分数段人数统计')
                plt.xlabel('分数段')
                plt.ylabel('人数')
                plt.xticks(rotation=45)
                plt.tight_layout()
                bar_file = f'bar_{class_name}_{subject}.png'
                plt.savefig(bar_file)
                plt.close()
                doc.add_picture(bar_file, width=Inches(6))
                
                # 5. 成绩趋势图
                plt.figure(figsize=(10, 6))
                plt.plot(range(len(class_data)), sorted(class_data[subject], reverse=True), 'b-')
                plt.axhline(y=mean_score, color='r', linestyle='--', label=f'平均分:{mean_score:.1f}')
                plt.title(f'{class_name}班级{subject}成绩趋势图')
                plt.xlabel('学生排名')
                plt.ylabel('分数')
                plt.legend()
                plt.tight_layout()
                trend_file = f'trend_{class_name}_{subject}.png'
                plt.savefig(trend_file)
                plt.close()
                doc.add_picture(trend_file, width=Inches(6))
                
                # 删除临时图片文件
                import os
                for img in [hist_file, box_file, pie_file, bar_file, trend_file]:
                    os.remove(img)
                    
            except Exception as e:
                print(f"生成{class_name}班级{subject}图表时出错: {str(e)}")
    
    # 添加班级间对比分析
    doc.add_heading('班级对比分析', level=1)
    
    for subject in subjects:
        doc.add_heading(f'{subject}科目班级对比', level=2)
        
        # 班级平均分对比条形图
        plt.figure(figsize=(10, 6))
        class_means = df.groupby('班级')[subject].mean()
        class_means.plot(kind='bar')
        plt.title(f'各班级{subject}平均分对比')
        plt.xlabel('班级')
        plt.ylabel('平均分')
        plt.tight_layout()
        plt.savefig(f'class_comparison_{subject}.png')
        plt.close()
        doc.add_picture(f'class_comparison_{subject}.png', width=Inches(6))
        os.remove(f'class_comparison_{subject}.png')
        
        # 班级各分数段比例对比堆积图
        plt.figure(figsize=(12, 6))
        class_data = []
        class_names = []
        for class_name, data in class_groups:
            class_names.append(class_name)
            stats = []
            for low, high, _ in score_ranges:
                percentage = len(data[(data[subject] >= low) & (data[subject] < high)]) / len(data) * 100
                stats.append(percentage)
            class_data.append(stats)
        
        class_data = np.array(class_data).T
        bottom = np.zeros(len(class_names))
        
        for i, (_, _, label) in enumerate(score_ranges):
            plt.bar(class_names, class_data[i], bottom=bottom, label=label)
            bottom += class_data[i]
        
        plt.title(f'各班级{subject}分数段比例对比')
        plt.xlabel('班级')
        plt.ylabel('比例(%)')
        plt.legend()
        plt.tight_layout()
        plt.savefig(f'score_distribution_comparison_{subject}.png')
        plt.close()
        doc.add_picture(f'score_distribution_comparison_{subject}.png', width=Inches(6))
        os.remove(f'score_distribution_comparison_{subject}.png')
    
    # 保存Word文档
    try:
        doc.save(word_file)
        print(f"分析报告已保存至 {word_file}")
    except Exception as e:
        print(f"保存Word文档时出错: {str(e)}")
    
    # 返回各班级各科目统计数据
    class_statistics = {}
    for class_name, class_data in class_groups:
        class_statistics[class_name] = {}
        for subject in subjects:
            score_distribution = []
            for low, high, label in score_ranges:
                count = len(class_data[(class_data[subject] >= low) & (class_data[subject] < high)])
                percentage = (count / len(class_data)) * 100
                score_distribution.append({
                    "分数段": label,
                    "人数": count,
                    "占比": percentage
                })
            
            class_statistics[class_name][subject] = {
                "基础统计": {
                    "平均分": class_data[subject].mean(),
                    "及格率": (class_data[subject] >= 60).mean() * 100,
                    "优秀率": (class_data[subject] >= 90).mean() * 100,
                    "差分率": (class_data[subject] < 60).mean() * 100,
                    "标准差": class_data[subject].std()
                },
                "分数段统计": score_distribution
            }
    
    return class_statistics

# 使用示例
if __name__ == "__main__":
    print("欢迎使用成绩分析工具！")
    print("易文胜制作，2024")
    print("------------------------------------------")
    #file_path = "成绩表.xlsx"  # Excel文件路径，需包含"班级"列和各科目成绩列
    file_path=input("请输入Excel文件全名：")
    try:
        results = analyze_scores_by_class(file_path)
        print("\n各班级各科目统计结果:")
        for class_name, subjects_stats in results.items():
            print(f"\n{class_name}班级:")
            for subject, stats in subjects_stats.items():
                print(f"\n{subject}:")
                print("基础统计:")
                for metric, value in stats["基础统计"].items():
                    print(f"{metric}: {value:.2f}")
                print("\n分数段统计:")
                for range_stat in stats["分数段统计"]:
                    print(f"{range_stat['分数段']}: {range_stat['人数']}人 ({range_stat['占比']:.2f}%)")
        print("\n\n分析报告已生成，请查看Word文档。")
    except Exception as e:
        print(f"程序执行出错: {str(e)}")